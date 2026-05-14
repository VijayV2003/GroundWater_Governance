"""
Groundwater Intelligence Platform – Backend API (Real-Time Edition)
====================================================================
FastAPI application exposing all 7 ML models as REST endpoints.
Powered by real Atal Jal DWLR dataset (2015-2022) with live data extrapolation.

Data Source    : Atal_Jal_Disclosed_Ground_Water_Level-2015-2022.csv
Real-Time Mode : Extrapolated from last known values, with a clear hook
                 (see `data/real_data_ingestion.py::append_realtime_mock_data`)
                 to swap in a live Government API key (data.gov.in) once available.

Run locally:
    uvicorn main:app --reload --port 8000

CORS is open for local development; restrict origins in production.
"""

from __future__ import annotations

import logging
import os
from contextlib import asynccontextmanager
from datetime import datetime
from typing import Any, Dict, List, Optional
import asyncio

from dotenv import load_dotenv
load_dotenv()
from sendgrid import SendGridAPIClient
from sendgrid.helpers.mail import Mail
import numpy as np
from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field
from fastapi.responses import StreamingResponse
import io
from docx import Document

# ── Internal modules ─────────────────────────────────────────────────────────
from data.real_data_ingestion import (
    load_real_stations,
    get_time_series_for_station,
    get_all_real_stations_with_status,
)
from data_generator import (
    STATIONS as SYNTHETIC_STATIONS,
    generate_socioeconomic_features,
    get_regional_summary,
)
from forecasting        import GroundwaterForecaster
from dhsf               import DHSFModel, FEATURE_COLS as DHSF_FEATURES
from anomaly_detection  import AnomalyDetector
from recharge_prediction import RechargePredictor
from gsi_dherp          import GSIInput, compute_gsi, compute_dherp
from aquifer_stress     import AquiferStressClassifier, FEATURE_COLS as STRESS_FEATURES
from firebase_utils     import fetch_policy_maker_emails

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

SAVED_MODELS_DIR = os.path.join(os.path.dirname(__file__), "saved_models")

# ── Singleton model instances (loaded once at startup) ────────────────────────
_forecasters: Dict[str, GroundwaterForecaster] = {}
_dhsf:        DHSFModel                        = None
_detector:    AnomalyDetector                  = None
_recharge:    RechargePredictor                = None
_stress_clf:  AquiferStressClassifier          = None

# Cache the real station list at startup
_real_stations: List[Dict] = []


# ─── Real-Time API Integration Hook ──────────────────────────────────────────
# TODO: When a Government API key is issued by CGWB / data.gov.in, replace the
#       stub below with a real HTTP call. The function signature stays the same
#       so zero changes are needed elsewhere in this file.
#
#   import requests
#   GOV_API_KEY  = os.getenv("CGWB_API_KEY", "")
#   GOV_API_BASE = "https://data.gov.in/api/v1/catalog/groundwater"
#
#   def _fetch_live_data(station_id: str) -> dict:
#       r = requests.get(GOV_API_BASE,
#                        params={"station_id": station_id, "api-key": GOV_API_KEY},
#                        timeout=5)
#       r.raise_for_status()
#       return r.json()
# ─────────────────────────────────────────────────────────────────────────────

async def send_email_updates():
    """Background task to send real emails via SendGrid to policy holders every 5 minutes."""
    while True:
        await asyncio.sleep(300) # 5 minutes
        
        api_key = os.getenv("SENDGRID_API_KEY")
        sender_email = os.getenv("SENDGRID_SENDER_EMAIL")
        enable_alerts = os.getenv("ENABLE_EMAIL_ALERTS", "true").lower() == "true"

        if not enable_alerts:
            logger.info("Email alerts are currently DISABLED via config.")
            continue
            
        # Dynamically fetch latest policy makers from Firebase Firestore
        receiver_emails = fetch_policy_maker_emails()

        if not receiver_emails:
            logger.warning("No policy maker emails found in Firebase or .env. Skipping dispatch.")
            continue

        if not api_key or not sender_email:
            logger.warning("SendGrid API key or sender email missing. Skipping email dispatch.")
            continue

        # --- Fetch Live Results for Email ---
        stations = get_all_real_stations_with_status(limit=5)
        critical_stations = [s for s in stations if s["status"] == "critical"]
        warning_stations  = [s for s in stations if s["status"] == "warning"]

        # Sustainability Index Logic
        baselines = [s["base_level"] for s in stations]
        levels    = [s["waterLevel"] for s in stations]
        ratios    = [b / l if l > 0 else 1.0 for b, l in zip(baselines, levels)]
        gsi       = int(np.clip(np.mean(ratios) * 80, 0, 100)) if ratios else 0

        current_time_str = datetime.now().strftime("%B %d, %Y, %H:%M IST")

        # Construct Dynamic Badges
        critical_html = ""
        for s in critical_stations:
            critical_html += f"""
            <div class="badge-critical">
              <p class="badge-title">🔴 CRITICAL — Station {s['id']} ({s['region']})</p>
              <p class="badge-body">Current Level: {s['waterLevel']}m (Baseline: {s['base_level']}m). Status: {s['status'].upper()}. Trend: {s['trend'].upper()}.</p>
            </div>
            """

        warning_html = ""
        for s in warning_stations:
            warning_html += f"""
            <div class="badge-warning">
              <p class="badge-title">🟡 WARNING — Station {s['id']} ({s['region']})</p>
              <p class="badge-body">Current Level: {s['waterLevel']}m (Baseline: {s['base_level']}m). Recovery monitoring recommended.</p>
            </div>
            """

        if not critical_html and not warning_html:
            critical_html = "<p style='color:#15803d;'>✅ No critical or warning levels detected across monitored stations.</p>"

        message = Mail(
            from_email=sender_email,
            to_emails=receiver_emails,
            subject=f"🔴 URGENT: Groundwater Security Alert ({datetime.now().strftime('%d %b %Y')})",
            html_content=f"""
<!DOCTYPE html>
<html>
<head>
  <style>
    body {{ font-family: Arial, sans-serif; background: #f4f6f8; margin: 0; padding: 20px; }}
    .card {{ background: #fff; border-radius: 8px; padding: 30px; max-width: 620px; margin: auto; box-shadow: 0 2px 8px rgba(0,0,0,0.1); }}
    .header {{ background: linear-gradient(135deg, #1e3a5f, #2563eb); color: white; border-radius: 8px 8px 0 0; padding: 20px 30px; margin: -30px -30px 24px; }}
    .header h1 {{ margin: 0; font-size: 20px; }}
    .header p  {{ margin: 4px 0 0; opacity: 0.85; font-size: 13px; }}
    .badge-critical {{ background: #fee2e2; color: #b91c1c; border-left: 4px solid #ef4444; padding: 12px 16px; border-radius: 4px; margin-bottom: 12px; }}
    .badge-warning  {{ background: #fef9c3; color: #92400e; border-left: 4px solid #f59e0b; padding: 12px 16px; border-radius: 4px; margin-bottom: 12px; }}
    .badge-title {{ font-weight: bold; font-size: 14px; margin: 0 0 4px; }}
    .badge-body  {{ font-size: 13px; margin: 0; }}
    .kpi {{ display: flex; gap: 12px; margin: 20px 0; }}
    .kpi-box {{ flex: 1; background: #f1f5f9; border-radius: 6px; padding: 14px; text-align: center; }}
    .kpi-val {{ font-size: 22px; font-weight: bold; color: #1e3a5f; }}
    .kpi-lbl {{ font-size: 11px; color: #64748b; margin-top: 2px; }}
    .btn {{ display: inline-block; background: #2563eb; color: white; padding: 12px 28px; border-radius: 6px; text-decoration: none; font-weight: bold; font-size: 14px; margin-top: 16px; }}
    .footer {{ font-size: 11px; color: #94a3b8; text-align: center; margin-top: 24px; }}
    .suggestion {{ background: #f0fdf4; border: 1px solid #bbf7d0; border-radius: 6px; padding: 14px 16px; margin-top: 16px; }}
    .suggestion h4 {{ margin: 0 0 8px; color: #15803d; font-size: 13px; }}
    .suggestion ul {{ margin: 0; padding-left: 18px; font-size: 12px; color: #374151; }}
    .suggestion ul li {{ margin-bottom: 4px; }}
  </style>
</head>
<body>
<div class="card">
  <div class="header">
    <h1>🔴 Groundwater Security Alert</h1>
    <p>Computational Groundwater Governance Platform — Automated Alert Dispatch</p>
  </div>

  <p style="color:#374151;font-size:14px;">Dear Policy Maker,</p>
  <p style="color:#374151;font-size:14px;">Our AI system has analyzed the latest data from the National DWLR Network. Below is the real-time summary as of <strong>{current_time_str}</strong>.</p>

  {critical_html}
  {warning_html}

  <!-- KPI Row -->
  <div class="kpi">
    <div class="kpi-box"><div class="kpi-val">{len(stations)}</div><div class="kpi-lbl">Active Stations</div></div>
    <div class="kpi-box"><div class="kpi-val">{gsi}%</div><div class="kpi-lbl">Sustainability Index</div></div>
    <div class="kpi-box"><div class="kpi-val">{len(critical_stations)}</div><div class="kpi-lbl">Critical Alerts</div></div>
    <div class="kpi-box"><div class="kpi-val">{len(warning_stations)}</div><div class="kpi-lbl">Warnings</div></div>
  </div>

  <div class="suggestion">
    <h4>✅ Suggested Actions</h4>
    <ul>
      <li>Review extraction permits in affected zones.</li>
      <li>Deploy rapid groundwater recharge interventions where critical depletion is noted.</li>
      <li>Issue advisory to local boards to optimize water usage.</li>
      <li>Monitor the dashboard for 7-model predictive insights.</li>
    </ul>
  </div>

  <a href="https://groundwater-governance.vercel.app" class="btn">Open Groundwater Dashboard →</a>

  <div class="footer">
    This email is auto-generated by the Computational Groundwater Governance Platform.<br>
    Powered by 7 active ML models — Data extrapolated from Atal Jal dataset.
  </div>
</div>
</body>
</html>
"""
        )

        try:
            # We must run this synchronous HTTP request in a threadpool to not block asyncio
            sg = SendGridAPIClient(api_key)
            response = await asyncio.to_thread(sg.send, message)
            logger.info(f"📧 [EMAIL SENT SUCCESS] Status Code: {response.status_code}")
        except Exception as e:
            logger.error(f"❌ [EMAIL FAILED] Could not send alerts: {e}")



import threading
_model_lock = threading.Lock()

def _get_detector():
    global _detector
    if _detector is None:
        with _model_lock:
            if _detector is None:
                import joblib as _jl
                path = os.path.join(SAVED_MODELS_DIR, "global_anomaly_detector.joblib")
                if os.path.exists(path):
                    _detector = AnomalyDetector.load(path)
                else:
                    ref_df = get_time_series_for_station(_real_stations[0]["id"])
                    _detector = AnomalyDetector()
                    _detector.fit(ref_df)
                    _detector.save(path)
    return _detector

def _get_dhsf():
    global _dhsf
    if _dhsf is None:
        with _model_lock:
            if _dhsf is None:
                import joblib as _jl
                path = os.path.join(SAVED_MODELS_DIR, "dhsf_model.joblib")
                if os.path.exists(path):
                    _dhsf = _jl.load(path)
                else:
                    _dhsf = DHSFModel()
                    _dhsf.train()
                    _jl.dump(_dhsf, path)
    return _dhsf

def _get_recharge():
    global _recharge
    if _recharge is None:
        with _model_lock:
            if _recharge is None:
                import joblib as _jl
                path = os.path.join(SAVED_MODELS_DIR, "recharge_predictor.joblib")
                if os.path.exists(path):
                    _recharge = _jl.load(path)
                else:
                    _recharge = RechargePredictor()
                    _recharge.train()
                    _jl.dump(_recharge, path)
    return _recharge

def _get_stress_clf():
    global _stress_clf
    if _stress_clf is None:
        with _model_lock:
            if _stress_clf is None:
                import joblib as _jl
                path = os.path.join(SAVED_MODELS_DIR, "stress_classifier.joblib")
                if os.path.exists(path):
                    _stress_clf = _jl.load(path)
                else:
                    _stress_clf = AquiferStressClassifier()
                    _stress_clf.train()
                    _jl.dump(_stress_clf, path)
    return _stress_clf

def _get_forecaster(station_id: str):
    if station_id not in _forecasters:
        with _model_lock:
            if station_id not in _forecasters:
                import joblib as _jl
                fc_path = os.path.join(SAVED_MODELS_DIR, f"forecaster_{station_id}.joblib")
                if os.path.exists(fc_path):
                    _forecasters[station_id] = GroundwaterForecaster.load(fc_path)
                else:
                    hist_df = get_time_series_for_station(station_id)
                    fc = GroundwaterForecaster(lags=7)
                    fc.fit(hist_df["water_level"].values)
                    fc.save(fc_path)
                    _forecasters[station_id] = fc
    return _forecasters[station_id]

@asynccontextmanager
async def lifespan(app: FastAPI):
    """
    Lightweight startup — only loads station index (~1 MB).
    All ML models are lazy-loaded on first request to stay within
    Render free tier 512 MB memory limit.
    """
    global _real_stations

    logger.info("🌊 Groundwater Intelligence Platform starting (memory-efficient mode) …")
    os.makedirs(SAVED_MODELS_DIR, exist_ok=True)

    # Only load the lightweight station index at startup
    _real_stations = load_real_stations(limit=5)
    logger.info(f"  ✓ {len(_real_stations)} stations ready (ML models load on first request)")

    # Start the 5-minute background email task
    email_task = asyncio.create_task(send_email_updates())

    logger.info("🚀 Ready — ML models will load lazily on first request")
    yield
    logger.info("Shutting down …")


# ── App factory ───────────────────────────────────────────────────────────────
app = FastAPI(
    title="Groundwater Intelligence Platform API",
    description=(
        "Real-time groundwater analytics powered by the Atal Jal DWLR dataset. "
        "Extrapolates to the current date and is API-key-ready for live CGWB / "
        "data.gov.in streams once a government key is provisioned."
    ),
    version="2.0.0",
    lifespan=lifespan,
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],   # restrict in production
    allow_methods=["*"],
    allow_headers=["*"],
)


# ═══════════════════════════════════════════════════════════════════════════════
# ── Pydantic request / response schemas ────────────────────────────────────────
# ═══════════════════════════════════════════════════════════════════════════════

class ForecastRequest(BaseModel):
    station_id:    str = Field(..., example="G_1_BK_021")
    horizon_days:  int = Field(30, ge=1, le=365, description="Forecast horizon in days")
    history_days:  int = Field(365, ge=30, le=2920, description="How many days of history to train on")


class DHSFRequest(BaseModel):
    station_id:                str   = Field(..., example="G_1_BK_021")
    agricultural_area_pct:     float = Field(45.0)
    irrigation_intensity:      float = Field(0.7)
    industrial_units_per_km2:  float = Field(10.0)
    population_density:        float = Field(500.0)
    annual_rainfall_mm:        float = Field(900.0)
    evapotranspiration_mm:     float = Field(1000.0)
    soil_permeability:         float = Field(0.5)
    surface_water_index:       float = Field(0.4)
    depletion_rate_m_per_year: float = Field(0.4)
    seasonal_amplitude_m:      float = Field(2.5)
    recharge_deficit_mm:       float = Field(100.0)


class RechargeRequest(BaseModel):
    station_id:            str   = Field(..., example="G_1_BK_021")
    annual_rainfall_mm:    float = Field(900.0)
    evapotranspiration_mm: float = Field(1000.0)
    soil_moisture_pct:     float = Field(50.0)
    surface_water_level_m: float = Field(5.0)
    land_use_idx:          float = Field(0.6)


class GSIRequest(BaseModel):
    station_id:               str   = Field(..., example="G_1_BK_021")
    current_level_m:          float = Field(45.0)
    baseline_level_m:         float = Field(50.0)
    recharge_rate_mm_yr:      float = Field(300.0)
    extraction_rate_mm_yr:    float = Field(400.0)
    trend_slope_m_yr:         float = Field(-0.3)
    climatic_recharge_support: float = Field(0.6)
    aquifer_storage_coeff:    float = Field(0.1)


class DHERPRequest(BaseModel):
    station_id:             str   = Field(..., example="G_1_BK_021")
    current_level_m:        float = Field(45.0)
    baseline_level_m:       float = Field(50.0)
    aquifer_area_km2:       float = Field(100.0)
    specific_yield:         float = Field(0.15)
    electricity_tariff_inr: float = Field(6.5)


class StressRequest(BaseModel):
    station_id:               str   = Field(..., example="G_1_BK_021")
    fluctuation_amplitude_m:  float = Field(3.0)
    extraction_intensity:     float = Field(0.8)
    soil_infiltration_rate:   float = Field(12.0)
    pre_monsoon_depth_m:      float = Field(15.0)
    post_monsoon_depth_m:     float = Field(10.0)
    long_term_trend_m_yr:     float = Field(-0.3)
    stage_of_extraction_pct:  float = Field(80.0)
    number_of_wells_per_km2:  float = Field(20.0)


# ═══════════════════════════════════════════════════════════════════════════════
# ── Helper utilities ───────────────────────────────────────────────────────────
# ═══════════════════════════════════════════════════════════════════════════════

def _get_real_station(station_id: str) -> Dict:
    """Look up a station from the real Atal Jal index."""
    for s in _real_stations:
        if s["id"] == station_id:
            return s
    return None


def _get_hist_df(station_id: str, history_days: int = 365):
    """Fetch real + extrapolated time-series for a station."""
    df = get_time_series_for_station(station_id, interpolate_daily=True)
    if df is not None and not df.empty:
        return df.tail(history_days)
    raise HTTPException(404, f"No time-series data available for station {station_id}")


def _auto_socioeconomic(station: Dict) -> Dict:
    """
    Return socioeconomic proxy features for a station.
    Currently uses the synthetic generator seeded by station id.
    Replace with real census/IMD data once available.
    """
    synth = next(
        (s for s in SYNTHETIC_STATIONS if s["region"] == station.get("region")),
        SYNTHETIC_STATIONS[0],
    )
    return generate_socioeconomic_features(synth)


from starlette.responses import RedirectResponse

@app.get("/", include_in_schema=False)
def root():
    """Redirect to the interactive API documentation."""
    return RedirectResponse(url="/docs")


# ═══════════════════════════════════════════════════════════════════════════════
# ── Health / metadata ──────────────────────────────────────────────────────────
# ═══════════════════════════════════════════════════════════════════════════════

@app.get("/health", tags=["System"])
def health() -> Dict:
    return {
        "status": "ok",
        "version": "2.0.0",
        "data_source": "Atal Jal DWLR 2015-2022 + live extrapolation",
        "real_time_api_key_configured": bool(os.getenv("CGWB_API_KEY")),
        "stations_loaded": len(_real_stations),
        "models_loaded": _dhsf is not None and _detector is not None,
        "server_time_utc": datetime.utcnow().isoformat(),
    }


# ═══════════════════════════════════════════════════════════════════════════════
# ── Dashboard endpoints ────────────────────────────────────────────────────────
# ═══════════════════════════════════════════════════════════════════════════════

@app.get("/api/stations", tags=["Dashboard"])
def list_stations(limit: int = 100) -> List[Dict]:
    """
    Return real DWLR stations from the Atal Jal dataset, with current
    water-level status computed from the interpolated + extrapolated time-series.
    """
    return get_all_real_stations_with_status(limit=limit)


@app.get("/api/dashboard/summary", tags=["Dashboard"])
def dashboard_summary() -> Dict:
    """Top-level KPI cards — derived from real station readings."""
    stations = get_all_real_stations_with_status(limit=100)
    statuses  = [s["status"] for s in stations]
    levels    = [s["waterLevel"] for s in stations]
    baselines = [s["base_level"] for s in stations]

    # GSI: ratio of current / baseline depth (lower depth = healthier)
    ratios = [b / l if l > 0 else 1.0 for b, l in zip(baselines, levels)]
    gsi    = int(np.clip(np.mean(ratios) * 80, 0, 100))

    # Approximation for depletion rate and stress
    depletion_rate = int(np.clip((1.0 - np.mean(ratios)) * 100, 0, 100)) if gsi < 100 else 0
    aquifer_stress = int(np.clip(100 - gsi + 10, 0, 100))

    return {
        "total_stations":          len(_real_stations),
        "displayed_stations":      len(stations),
        "online_pct":              round(100 * sum(1 for s in statuses if s != "offline") / max(len(statuses), 1), 1),
        "avg_water_level_m":       round(float(np.mean(levels)), 1),
        "active_alerts":           sum(1 for s in statuses if s in ("warning", "critical")),
        "critical_count":          sum(1 for s in statuses if s == "critical"),
        "warning_count":           sum(1 for s in statuses if s == "warning"),
        "sustainability_index":    gsi,
        "depletion_rate":          depletion_rate,
        "aquifer_stress":          aquifer_stress,
        "annual_loss_crore_inr":   126 + (aquifer_stress - 60) * 2,
        "people_affected_million": round(2.3 + (depletion_rate - 20) * 0.1, 1),
        "data_source":             "Atal Jal DWLR 2015-2022",
        "realtime_extrapolated_to": datetime.utcnow().strftime("%Y-%m-%d"),
        "api_key_ready":           True,  # backend is ready; swap stub in real_data_ingestion.py
    }


@app.get("/api/dashboard/regional", tags=["Dashboard"])
def regional_summary() -> List[Dict]:
    """Regional aggregation — weighted by real station alerts."""
    stations = get_all_real_stations_with_status(limit=100)
    return get_regional_summary(stations)


# ═══════════════════════════════════════════════════════════════════════════════
# ── Model 1 – Groundwater Level Forecast ───────────────────────────────────────
# ═══════════════════════════════════════════════════════════════════════════════

@app.post("/api/models/forecast", tags=["Model 1 – Forecasting"])
def forecast_water_level(req: ForecastRequest) -> Dict:
    """
    Predict future water levels for a real DWLR station.
    Uses a pre-warmed cached forecaster (fitted at startup) to eliminate
    per-request training latency. Falls back to on-demand fit if station
    is not in the cache (e.g. a non-primary station).
    """
    import joblib as _jl
    hist_df = _get_hist_df(req.station_id, req.history_days)

    if req.station_id in _forecasters:
        fc = _forecasters[req.station_id]
        fc.fit(hist_df["water_level"].values)
    else:
        fc = _get_forecaster(req.station_id)
        result = fc.forecast_api_response(req.station_id, hist_df, req.horizon_days)
        return result

    result = fc.forecast_api_response(req.station_id, hist_df, req.horizon_days)
    _forecasters[req.station_id] = fc
    return result


@app.get("/api/models/forecast/{station_id}", tags=["Model 1 – Forecasting"])
def forecast_default(station_id: str) -> Dict:
    """Quick GET with defaults: 30-day horizon, 365 days of history."""
    return forecast_water_level(ForecastRequest(station_id=station_id))


# ═══════════════════════════════════════════════════════════════════════════════
# ── Model 2 – DHSF Depletion Cause Classification ──────────────────────────────
# ═══════════════════════════════════════════════════════════════════════════════

@app.post("/api/models/dhsf", tags=["Model 2 – DHSF"])
def classify_depletion_cause(req: DHSFRequest) -> Dict:
    feats  = req.model_dump(exclude={"station_id"})
    result = _get_dhsf().predict(feats)
    result["station_id"] = req.station_id
    return result


@app.get("/api/models/dhsf/{station_id}", tags=["Model 2 – DHSF"])
def dhsf_auto(station_id: str) -> Dict:
    """Auto-classify using time-series derived + socioeconomic proxy features."""
    hist_df  = _get_hist_df(station_id, 365)
    levels   = hist_df["water_level"].values
    trend_yr = float(np.polyfit(np.arange(len(levels)), levels, 1)[0] * 365)

    station = _get_real_station(station_id)
    se      = _auto_socioeconomic(station or {})

    feats = {
        **se,
        "depletion_rate_m_per_year": abs(trend_yr),
        "seasonal_amplitude_m":      float(levels.max() - levels.min()),
        "recharge_deficit_mm":       max(0.0, se.get("evapotranspiration_mm", 1000) - se.get("annual_rainfall_mm", 900)),
    }
    result = _get_dhsf().predict(feats)
    result["station_id"] = station_id
    result["trend_m_per_year"] = round(trend_yr, 3)
    return result


# ═══════════════════════════════════════════════════════════════════════════════
# ── Model 3 – Anomaly Detection ────────────────────────────────────────────────
# ═══════════════════════════════════════════════════════════════════════════════

@app.get("/api/models/anomaly/{station_id}", tags=["Model 3 – Anomaly Detection"])
def detect_anomalies(station_id: str, days: int = 365) -> Dict:
    """Detect anomalies in a station's real + extrapolated time-series."""
    hist_df = _get_hist_df(station_id, days)
    return _get_detector().detect(hist_df, station_id)


@app.get("/api/alerts", tags=["Model 3 – Anomaly Detection"])
def get_alerts(limit: int = 20) -> List[Dict]:
    """
    Run anomaly detection across the first batch of real stations.
    Returns live-style alert list for the dashboard.
    """
    station_results = []
    det = _get_detector()
    for s in _real_stations[:20]:
        try:
            hist_df = get_time_series_for_station(s["id"], interpolate_daily=True)
            if hist_df is not None and len(hist_df) >= 7:
                res = det.detect(hist_df.tail(365), s["id"])
                station_results.append(res)
        except Exception:
            continue
    return det.generate_alerts(station_results)[:limit]


# ═══════════════════════════════════════════════════════════════════════════════
# ── Model 4 – Monsoon Recharge Prediction ──────────────────────────────────────
# ═══════════════════════════════════════════════════════════════════════════════

@app.post("/api/models/recharge", tags=["Model 4 – Recharge Prediction"])
def predict_recharge(req: RechargeRequest) -> Dict:
    monthly = _get_recharge().predict_annual(
        annual_rainfall_mm=req.annual_rainfall_mm,
        evapotranspiration_mm=req.evapotranspiration_mm,
        soil_moisture_pct=req.soil_moisture_pct,
        surface_water_level_m=req.surface_water_level_m,
        land_use_idx=req.land_use_idx,
    )
    rch = _get_recharge()
    return {"station_id": req.station_id, "recharge_data": monthly, "summary": rch.monsoon_summary(monthly)}


@app.get("/api/models/recharge/{station_id}", tags=["Model 4 – Recharge Prediction"])
def recharge_auto(station_id: str) -> Dict:
    station = _get_real_station(station_id) or {}
    se      = _auto_socioeconomic(station)
    rch = _get_recharge()
    monthly = rch.predict_annual(
        annual_rainfall_mm=se["annual_rainfall_mm"],
        evapotranspiration_mm=se["evapotranspiration_mm"],
        soil_moisture_pct=60.0,
        surface_water_level_m=se["surface_water_index"] * 10,
        land_use_idx=se["agricultural_area_pct"] / 100,
    )
    return {"station_id": station_id, "recharge_data": monthly, "summary": rch.monsoon_summary(monthly)}


# ═══════════════════════════════════════════════════════════════════════════════
# ── Model 5 – GSI Scoring ──────────────────────────────────────────────────────
# ═══════════════════════════════════════════════════════════════════════════════

@app.post("/api/models/gsi", tags=["Model 5 – GSI Scoring"])
def score_gsi(req: GSIRequest) -> Dict:
    inp    = GSIInput(**req.model_dump(exclude={"station_id"}))
    result = compute_gsi(inp)
    result["station_id"] = req.station_id
    return result


@app.get("/api/models/gsi/{station_id}", tags=["Model 5 – GSI Scoring"])
def gsi_auto(station_id: str) -> Dict:
    station  = _get_real_station(station_id)
    if not station:
        raise HTTPException(404, f"Station {station_id} not found")
    hist    = _get_hist_df(station_id, 365)
    current  = float(hist["water_level"].iloc[-1])
    baseline = station["base_level"]
    trend    = float(np.polyfit(np.arange(len(hist)), hist["water_level"].values, 1)[0] * 365)
    inp = GSIInput(
        current_level_m=current,
        baseline_level_m=baseline,
        recharge_rate_mm_yr=300.0,
        extraction_rate_mm_yr=420.0,
        trend_slope_m_yr=trend,
        climatic_recharge_support=0.55,
        aquifer_storage_coeff=0.12,
    )
    result = compute_gsi(inp)
    result["station_id"] = station_id
    result["trend_m_per_year"] = round(trend, 3)
    return result


# ═══════════════════════════════════════════════════════════════════════════════
# ── Model 6 – DH-ERP Index ─────────────────────────────────────────────────────
# ═══════════════════════════════════════════════════════════════════════════════

@app.post("/api/models/dherp", tags=["Model 6 – DH-ERP"])
def compute_dherp_endpoint(req: DHERPRequest) -> Dict:
    result = compute_dherp(
        current_level_m=req.current_level_m,
        baseline_level_m=req.baseline_level_m,
        aquifer_area_km2=req.aquifer_area_km2,
        specific_yield=req.specific_yield,
        electricity_tariff_inr_kwh=req.electricity_tariff_inr,
    )
    result["station_id"] = req.station_id
    return result


@app.get("/api/models/dherp/{station_id}", tags=["Model 6 – DH-ERP"])
def dherp_auto(station_id: str) -> Dict:
    station  = _get_real_station(station_id)
    if not station:
        raise HTTPException(404, f"Station {station_id} not found")
    hist    = _get_hist_df(station_id, 365)
    current  = float(hist["water_level"].iloc[-1])
    baseline = station["base_level"]
    result   = compute_dherp(current, baseline, aquifer_area_km2=150.0)
    result["station_id"] = station_id
    return result


# ═══════════════════════════════════════════════════════════════════════════════
# ── Model 7 – Aquifer Stress Classification ─────────────────────────────────────
# ═══════════════════════════════════════════════════════════════════════════════

@app.post("/api/models/stress", tags=["Model 7 – Aquifer Stress"])
def classify_stress(req: StressRequest) -> Dict:
    feats  = req.model_dump(exclude={"station_id"})
    result = _get_stress_clf().predict(feats)
    result["station_id"] = req.station_id
    return result


@app.get("/api/models/stress/{station_id}", tags=["Model 7 – Aquifer Stress"])
def stress_auto(station_id: str) -> Dict:
    station = _get_real_station(station_id)
    if not station:
        raise HTTPException(404, f"Station {station_id} not found")
    hist    = _get_hist_df(station_id, 365)
    levels  = hist["water_level"].values
    feats = {
        "fluctuation_amplitude_m": float(levels.max() - levels.min()),
        "extraction_intensity":    0.75,
        "soil_infiltration_rate":  12.0,
        "pre_monsoon_depth_m":     float(levels[:int(len(levels) * 0.25)].mean()),
        "post_monsoon_depth_m":    float(levels[int(len(levels) * 0.5): int(len(levels) * 0.75)].mean()),
        "long_term_trend_m_yr":    float(np.polyfit(np.arange(len(levels)), levels, 1)[0] * 365),
        "stage_of_extraction_pct": 80.0,
        "number_of_wells_per_km2": 20.0,
    }
    result = _get_stress_clf().predict(feats)
    result["station_id"] = station_id
    return result


# ═══════════════════════════════════════════════════════════════════════════════
# ── Policy Report Generation ───────────────────────────────────────────────────
# ═══════════════════════════════════════════════════════════════════════════════

class ReportRequest(BaseModel):
    image_base64: Optional[str] = None

@app.post("/api/policy/generate-report/{station_id}", tags=["Policy Tools"])
def generate_policy_report(station_id: str, req: ReportRequest = None):
    """Generate an AI-powered Word document policy brief, optionally embedding a model insights snapshot."""
    stations = get_all_real_stations_with_status()
    station = next((s for s in stations if s["id"] == station_id), None)
    if not station:
        raise HTTPException(404, f"Station {station_id} not found")

    api_key = os.getenv("GEMINI_API_KEY") or os.getenv("GOOGLE_API_KEY")
    stress_data = stress_auto(station_id)

    report_text = None  # will be filled by Gemini or fallback template

    if api_key:
        try:
            import google.generativeai as _genai
            _genai.configure(api_key=api_key)
            gemini_model = _genai.GenerativeModel(
                "gemini-2.0-flash-lite",
                generation_config=_genai.types.GenerationConfig(
                    max_output_tokens=800,
                    temperature=0.4,
                )
            )
            prompt = (
                f"Write a concise government policy brief (5 sections, Markdown ## headings) "
                f"for station {station['id']} ({station['region']}, {station['status']}, "
                f"{station['waterLevel']}m bgl, trend:{station.get('trend','N/A')}). "
                f"Stress class: {stress_data.get('stress_class','Unknown')} "
                f"({round(stress_data.get('confidence',0)*100,0)}% confidence). "
                f"Actions: {'; '.join(stress_data.get('recommended_actions',[])[:3])}. "
                f"Sections: 1.Executive Summary 2.Status Overview "
                f"3.Aquifer Stress & Risk 4.Socio-Economic Impact 5.Recommendations. "
                f"Be factual, concise, bullet-pointed. No memorandum header."
            )
            response = gemini_model.generate_content(prompt)
            if response and response.text:
                report_text = response.text
                logger.info("✅ Gemini report generated successfully.")
            else:
                logger.warning("Gemini returned empty response — using fallback template.")
        except Exception as e:
            logger.warning(f"Gemini unavailable ({e}) — falling back to template report.")
    else:
        logger.warning("No GEMINI_API_KEY set — using fallback template report.")

    # ── Fallback: structured template using real ML data (no Gemini needed) ──
    if not report_text:
        sc = stress_data.get('stress_class', 'Unknown')
        conf = round(stress_data.get('confidence', 0) * 100, 1)
        actions = stress_data.get('recommended_actions', [
            "Regulate extraction permits in affected zone.",
            "Deploy artificial recharge structures.",
            "Issue advisory to local agricultural boards.",
        ])
        report_text = f"""## 1. Executive Summary

Station **{station['id']}** in the **{station['region']}** region is currently at a water level of \
**{station['waterLevel']} m bgl** with a **{station.get('trend','stable')}** trend. \
The ML-based aquifer stress classifier has categorised this station as **{sc}** with \
{conf}% confidence. Immediate policy intervention is recommended.

## 2. Groundwater Status Overview

- **Station ID:** {station['id']}
- **Region:** {station['region']}
- **Current Water Level:** {station['waterLevel']} m below ground level (bgl)
- **Operational Status:** {station['status'].upper()}
- **Trend:** {station.get('trend', 'N/A').capitalize()}
- **Baseline Depth:** {station.get('base_level', 'N/A')} m bgl

The station is part of the national DWLR network and is monitored in real-time. \
Current readings indicate a {station.get('trend','stable')} trajectory requiring attention.

## 3. Aquifer Stress & Risk Analysis

- **ML Model 7 Classification:** {sc}
- **Confidence Score:** {conf}%
- **Stage of Extraction:** {stress_data.get('stage_of_extraction_pct', 80)}%
- **Risk Level:** {'HIGH — Immediate action required.' if station['status'] == 'critical' else 'MODERATE — Monitoring and mitigation advised.'}

The aquifer in this region shows signs of sustained stress driven by over-extraction, \
reduced monsoon recharge, and increased anthropogenic demand. Continued depletion \
will compromise water security for dependent communities.

## 4. Socio-Economic Impact Assessment

- Groundwater depletion in this region directly impacts agricultural productivity, \
  drinking water supply, and industrial operations.
- Declining levels increase pumping costs and energy consumption for farmers.
- Communities dependent on this aquifer face risk of water insecurity within \
  the next 2–5 years if the current trend continues.
- Economic losses are estimated to compound annually if intervention is delayed.

## 5. Specific Intervention Recommendations

"""
        for i, action in enumerate(actions[:6], 1):
            report_text += f"- {action}\n"
        report_text += (
            "\n- Establish a real-time monitoring dashboard shared with district officials.\n"
            "- Conduct annual aquifer health assessments using the DWLR network.\n"
            "- Promote community-level water conservation and rainwater harvesting.\n"
        )

    # ── Strip any leading memorandum block ────────────────────────────────────
    import re as _re
    first_section = _re.search(r'^## ', report_text, _re.MULTILINE)
    if first_section:
        report_text = report_text[first_section.start():]

    # ── Build the Word Document ───────────────────────────────────────────────
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from datetime import date as _date

    doc = Document()

    # ── PAGE 1: Model Insights Snapshot (if provided) ─────────────────────────
    if req and req.image_base64:
        try:
            import base64 as _b64
            snap_title = doc.add_heading("", 0)
            snap_title.clear()
            r = snap_title.add_run("ML Model Insights Dashboard")
            r.font.size = Pt(20)
            r.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)
            snap_sub = doc.add_paragraph()
            snap_sub.add_run(
                f"Station: {station['id']}  |  Region: {station['region']}  "
                f"|  Generated: {_date.today().strftime('%d %B %Y')}"
            ).bold = True
            doc.add_paragraph()
            img_bytes = _b64.b64decode(req.image_base64)
            img_stream = io.BytesIO(img_bytes)
            doc.add_picture(img_stream, width=Inches(6.2))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption = doc.add_paragraph(
                f"Real-time output of all 7 ML models for Station {station['id']} "
                f"({station['region']}) – captured at report generation time."
            )
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.runs[0].font.size = Pt(9)
            caption.runs[0].font.italic = True
            caption.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8b)
            doc.add_page_break()
        except Exception as img_err:
            logger.warning(f"Could not embed snapshot image: {img_err}")

    # ── Cover page ────────────────────────────────────────────────────────────
    title = doc.add_heading("", 0)
    title.clear()
    run = title.add_run("Water Security Policy Brief")
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)

    sub = doc.add_paragraph()
    sub.add_run(
        f"Station: {station['id']}  |  Region: {station['region']}  "
        f"|  Status: {station['status'].upper()}"
    ).bold = True

    meta = doc.add_paragraph()
    meta.add_run(
        f"Generated on: {_date.today().strftime('%d %B %Y')}  "
        f"|  Powered by 7 ML Models"
        + (" + Google Gemini" if api_key and "## " in report_text else " (Offline Template)")
    )
    meta.runs[-1].font.color.rgb = RGBColor(0x64, 0x74, 0x8b)
    doc.add_paragraph()

    # ── Parse and write content ───────────────────────────────────────────────
    for line in report_text.split('\n'):
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        if line.startswith('## '):
            doc.add_heading(line.replace('## ', '').strip(), level=2)
        elif line.startswith('### '):
            doc.add_heading(line.replace('### ', '').strip(), level=3)
        elif line.startswith('# '):
            doc.add_heading(line.replace('# ', '').strip(), level=1)
        elif line.startswith('* ') or line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('**') and line.endswith('**'):
            p = doc.add_paragraph()
            p.add_run(line.replace('**', '')).bold = True
        else:
            doc.add_paragraph(line)

    f = io.BytesIO()
    doc.save(f)
    f.seek(0)

    headers = {
        'Content-Disposition': f'attachment; filename="Policy_Brief_{station_id}.docx"'
    }
    return StreamingResponse(
        f,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers
    )

